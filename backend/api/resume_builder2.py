import requests
from sentence_transformers import SentenceTransformer, util
from docx import Document
from datetime import datetime

# Login for Auth
# loginURL = "http://127.0.0.1:8000/api/login/"
# loginCred = {
#     "email": "htruong@crimson.ua.edu",
#     "password": "Yay123456789!"
# }
# mockLoginResponse = requests.post(loginURL, json=loginCred)

# # Auth
# if mockLoginResponse.status_code == 200:
#     access_token = mockLoginResponse.json().get('access')  # Safely access 'access' key
# else:
#     print(f"Login failed: {mockLoginResponse.status_code}, {mockLoginResponse.text}")

def generate_resume(access_token, position):
    # Get All Experiences
    getExperiencesURL = "http://127.0.0.1:8000/api/experiences/"
    getProfilesURL = "http://127.0.0.1:8000/api/profiles/"
    # Get All Profile
    # Pass access token into header of the API call
    headers = {
        "Authorization": f"Bearer {access_token}"
    }
    getExperiencesResponse = requests.get(getExperiencesURL, headers=headers)
    getProfilesResponse = requests.get(getProfilesURL, headers=headers)
    # Extract the returned data
    rawData = []
    if getExperiencesResponse.status_code == 200:
        rawData = getExperiencesResponse.json()
        # print(rawData)
    else:
        print(f"Error: {getExperiencesResponse.status_code}, {getExperiencesResponse.text}")

    if getProfilesResponse.status_code == 200:
        profile = getProfilesResponse.json()[0]
    else:
        print(f"Error: {getExperiencesResponse.status_code}, {getExperiencesResponse.text}")

    # Loop through the experiences json objects and transform the data into a list of objects with structure:
        # exp_id: id of the experience
        # exp_data: parse the experience json into text
    transformed = []
    for exp in rawData:
        # Ensure exp is a dictionary and we access it using the key
        i = {
            "exp_id": exp.get('id'),
            "exp_data": f"{exp.get('job_title', 'N/A')} - {exp.get('description', 'N/A')}",
            "similarity_score": 0
        }
        transformed.append(i)
    # print(transformed)

    # Compare similarity between job_posting and each of your experiences
    job_posting = """Are you passionate about building scalable, high-performance applications? We’re looking for a Software Engineer to join our backend development team and help us design, develop, and optimize our cloud-based platforms. In this role, you’ll work with a talented group of engineers to build robust APIs, improve system architecture, and implement best coding practices to ensure seamless user experiences.

    Responsibilities:
    - Develop and maintain backend services, APIs, and databases to support our web and mobile applications.
    - Collaborate with frontend engineers, designers, and product managers to deliver high-quality features.
    - Optimize application performance, scalability, and security.
    - Write clean, efficient, and well-documented code following industry best practices.
    - Debug, troubleshoot, and resolve technical issues as they arise.

    Requirements:
    - Proficiency in Python, Node.js, or Java with experience in backend frameworks such as Django, Express.js, or Spring Boot.
    - Strong understanding of RESTful APIs, microservices, and database design (SQL & NoSQL).
    - Experience working with cloud platforms like AWS, GCP, or Azure.
    - Knowledge of authentication protocols (OAuth, JWT) and security best practices.
    - Strong problem-solving skills and ability to work both independently and in a team environment.

    If you’re excited to work on cutting-edge technology and make an impact in a fast-growing company, apply today! 🚀"""


    # MODEL FOR RELEVANT RANKING OF THE EXPERIENCES TO THE JOB POSTING
    model = SentenceTransformer("sentence-transformers/msmarco-MiniLM-L6-cos-v5")

    # Encode job posting and experiences
    job_embedding = model.encode(job_posting, convert_to_tensor=True)
    experience_embeddings = model.encode([exp["exp_data"] for exp in transformed], convert_to_tensor=True)

    # Compute similarity scores
    similarity_scores = util.pytorch_cos_sim(job_embedding, experience_embeddings)[0]

    # Update experience data with similarity scores
    for i, exp in enumerate(transformed):
        exp["similarity_score"] = similarity_scores[i].item()  # Convert tensor to Python float

    count = 3 #enter the number of top relevant jobs you would like to include
    # Sort experiences by similarity score (descending)
    top_relevant = sorted(transformed, key=lambda x: x["similarity_score"], reverse=True)[:count]
    shortlisted = []
    #Loop through the list of top relevant experiences
    for exp in top_relevant:
        # filter the rawData experience objects to only contain the top relevant
        for i in rawData:
            if i.get('id') == exp.get('exp_id'):
                shortlisted.append(i)
    # sort the experiences by date
    sorted_by_date = sorted(shortlisted, key=lambda x: x["from_date"], reverse=True)


    #method to write experience into the doc
    from datetime import datetime

    def addExperienceToDoc(expObj, doc):
        # Check if 'from_date' exists and is a string in the expected format
        try:
            from_date = datetime.strptime(expObj.get('from_date'), "%Y-%m-%dT%H:%M:%SZ")
        except (ValueError, TypeError):
            from_date = None  # If from_date is invalid or not present, set it to None
        
        # Check if 'to_date' exists and is a string in the expected format
        try:
            to_date = datetime.strptime(expObj.get('to_date'), "%Y-%m-%dT%H:%M:%SZ") if expObj.get('to_date') else None
        except (ValueError, TypeError):
            to_date = None  # If to_date is invalid or not present, set it to None
        
        doc.add_heading(f"{expObj.get('job_title')} | {expObj.get('company')}", level=3)
        expSection = doc.add_paragraph()

        # Ensure from_date and to_date are properly handled
        if from_date:
            expSection.add_run(f"{from_date.strftime('%m/%d/%Y')} – ")
        else:
            expSection.add_run("N/A – ")

        if to_date:
            expSection.add_run(f"{to_date.strftime('%m/%d/%Y')}\n")
        else:
            expSection.add_run("Present\n")
        
        expSection.add_run(f"{expObj.get('location')}\n").italic = True
        expSection.add_run("Description: ").bold = True
        expSection.add_run(f"{expObj.get('description')}")


    doc = Document()
    # add heading
    # Title (Name)
    doc.add_heading(f"{profile.get("full_name")}",level=1)
    # Contact Information (Email, phone, linkedIn)
    doc.add_paragraph(f"Email: {profile.get("email")} | Phone: {profile.get("phone")}\nLinkedIn: {profile.get("linkedin")}\nWebsite: {profile.get("website")}")
    # Eduction Section
    doc.add_heading('Education',level=2)
    edu = doc.add_paragraph()
    edu.add_run(f"{profile.get("latest_edu_name")}\n").bold = True
    edu.add_run(f"{profile.get("latest_edu_from_date").strftime('%m/%d/%Y')} - {profile.get("latest_edu_to_date").strftime('%m/%d/%Y')}").italic = True
    edu.add_run(f"{profile.get("latest_edu_desc")}\n")
    # Working Experience Section
    doc.add_heading('Working Experiences', level=2)
    for i in sorted_by_date:
        addExperienceToDoc(i, doc)
    # Skills Section
    doc.add_heading('Skills', level=2)
    skill = doc.add_paragraph()
    for i in sorted_by_date:
        skillList = i.get('skills')
        skill.add_run(f"{skillList[0]}")
        for x in range(1,len(skillList)):
            skill.add_run(f", {skillList[x]}")

    doc.save("resume.docx")
    return doc